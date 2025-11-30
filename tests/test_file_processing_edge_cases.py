# tests/test_file_processing_edge_cases.py
"""
Edge case tests for file processing across all processors.
Tests unusual files, corrupted content, encoding issues, and boundary conditions.
"""

import pytest
from pathlib import Path
from unittest.mock import Mock
import openpyxl
from openpyxl.styles import Font
from docx import Document

from yakulingo.processors.excel_processor import ExcelProcessor
from yakulingo.processors.word_processor import WordProcessor
from yakulingo.processors.pptx_processor import PptxProcessor
from yakulingo.processors.pdf_processor import PdfProcessor
from yakulingo.models.types import FileType, TextBlock


# =============================================================================
# Tests: Empty Files
# =============================================================================

class TestEmptyFiles:
    """Tests for processing empty files"""

    def test_empty_excel_file(self, tmp_path):
        """Process empty Excel file"""
        file_path = tmp_path / "empty.xlsx"
        wb = openpyxl.Workbook()
        wb.save(file_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        assert blocks == []

    def test_excel_with_only_empty_cells(self, tmp_path):
        """Process Excel file with only empty cells"""
        file_path = tmp_path / "empty_cells.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = ""
        ws['A2'] = "   "  # Whitespace only
        ws['A3'] = None
        wb.save(file_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        assert len(blocks) == 0  # Empty/whitespace cells should be skipped

    def test_empty_word_file(self, tmp_path):
        """Process empty Word file"""
        file_path = tmp_path / "empty.docx"
        doc = Document()
        doc.save(file_path)

        processor = WordProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        assert blocks == []

    def test_word_with_only_empty_paragraphs(self, tmp_path):
        """Process Word file with only empty paragraphs"""
        file_path = tmp_path / "empty_paragraphs.docx"
        doc = Document()
        doc.add_paragraph("")
        doc.add_paragraph("   ")
        doc.add_paragraph("\n\t")
        doc.save(file_path)

        processor = WordProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        # Empty paragraphs should be skipped
        assert len(blocks) == 0


# =============================================================================
# Tests: Large Files
# =============================================================================

class TestLargeFiles:
    """Tests for processing large files"""

    def test_excel_many_rows(self, tmp_path):
        """Process Excel file with many rows"""
        file_path = tmp_path / "many_rows.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active

        # Add 1000 rows
        for i in range(1, 1001):
            ws[f'A{i}'] = f'Row {i} text'

        wb.save(file_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        assert len(blocks) == 1000

    def test_excel_many_columns(self, tmp_path):
        """Process Excel file with many columns"""
        file_path = tmp_path / "many_columns.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active

        # Add 100 columns
        for i in range(1, 101):
            ws.cell(row=1, column=i).value = f'Column {i}'

        wb.save(file_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        assert len(blocks) == 100

    def test_excel_many_sheets(self, tmp_path):
        """Process Excel file with many sheets"""
        file_path = tmp_path / "many_sheets.xlsx"
        wb = openpyxl.Workbook()

        # Add 20 sheets with content
        for i in range(20):
            if i == 0:
                ws = wb.active
                ws.title = f"Sheet{i}"
            else:
                ws = wb.create_sheet(f"Sheet{i}")
            ws['A1'] = f'Content on Sheet {i}'

        wb.save(file_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        assert len(blocks) == 20

    def test_word_many_paragraphs(self, tmp_path):
        """Process Word file with many paragraphs"""
        file_path = tmp_path / "many_paragraphs.docx"
        doc = Document()

        for i in range(500):
            doc.add_paragraph(f'Paragraph {i} with text')

        doc.save(file_path)

        processor = WordProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        assert len(blocks) == 500


# =============================================================================
# Tests: Unicode and Special Characters
# =============================================================================

class TestUnicodeHandling:
    """Tests for Unicode and special character handling"""

    def test_excel_with_emoji(self, tmp_path):
        """Process Excel file with emoji"""
        file_path = tmp_path / "emoji.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = '絵文字テスト 🎌🗾'
        ws['A2'] = '複合絵文字 👨‍👩‍👧‍👦'
        wb.save(file_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        assert len(blocks) == 2
        assert '🎌' in blocks[0].text

    def test_excel_with_cjk_characters(self, tmp_path):
        """Process Excel file with CJK characters"""
        file_path = tmp_path / "cjk.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = '日本語テキスト'  # Japanese
        ws['A2'] = '中文文本'        # Chinese
        ws['A3'] = '한국어 텍스트'    # Korean
        wb.save(file_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        assert len(blocks) == 3

    def test_excel_with_rtl_text(self, tmp_path):
        """Process Excel file with right-to-left text"""
        file_path = tmp_path / "rtl.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = 'مرحبا بالعالم'  # Arabic
        ws['A2'] = 'שלום עולם'      # Hebrew
        wb.save(file_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        assert len(blocks) == 2

    def test_excel_with_zero_width_characters(self, tmp_path):
        """Process Excel file with zero-width characters"""
        file_path = tmp_path / "zero_width.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        # Zero-width space, joiner, non-joiner
        ws['A1'] = 'テスト\u200b\u200c\u200dテキスト'
        wb.save(file_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        assert len(blocks) == 1

    def test_word_with_special_punctuation(self, tmp_path):
        """Process Word file with special punctuation"""
        file_path = tmp_path / "punctuation.docx"
        doc = Document()
        doc.add_paragraph('「日本語の引用符」')
        doc.add_paragraph('【括弧のテスト】')
        doc.add_paragraph('～波ダッシュ～')
        doc.save(file_path)

        processor = WordProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        assert len(blocks) == 3


# =============================================================================
# Tests: Mixed Content
# =============================================================================

class TestMixedContent:
    """Tests for files with mixed content types"""

    def test_excel_with_formulas(self, tmp_path):
        """Process Excel file with formulas"""
        file_path = tmp_path / "formulas.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = 100
        ws['A2'] = 200
        ws['A3'] = '=SUM(A1:A2)'  # Formula
        ws['B1'] = 'テキストセル'
        wb.save(file_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        # Should include text cell, may or may not include formula
        text_blocks = [b for b in blocks if 'テキスト' in b.text]
        assert len(text_blocks) >= 1

    def test_excel_with_dates(self, tmp_path):
        """Process Excel file with dates"""
        from datetime import datetime
        file_path = tmp_path / "dates.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = datetime(2024, 1, 15)
        ws['A2'] = '2024年1月15日'  # Text date
        ws['B1'] = 'Normal text'
        wb.save(file_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        # Should extract text cells
        assert len(blocks) >= 1

    def test_excel_with_numbers_and_text(self, tmp_path):
        """Process Excel file with mixed numbers and text"""
        file_path = tmp_path / "mixed.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = 12345  # Number only - should skip
        ws['A2'] = '12345円'  # Number with unit - should translate
        ws['A3'] = 'Item 123'  # Text with number - should translate
        ws['A4'] = 'Pure text'  # Pure text - should translate
        wb.save(file_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        # Number-only should be skipped, others included
        assert len(blocks) >= 2


# =============================================================================
# Tests: Font Handling
# =============================================================================

class TestFontHandling:
    """Tests for font extraction and handling"""

    def test_excel_preserves_font_info(self, tmp_path):
        """ExcelProcessor extracts font information"""
        file_path = tmp_path / "fonts.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = 'Mincho text'
        ws['A1'].font = Font(name='MS Mincho', size=12)
        ws['A2'] = 'Gothic text'
        ws['A2'].font = Font(name='MS Gothic', size=14)
        wb.save(file_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        # Verify blocks have metadata
        for block in blocks:
            assert block.metadata is not None or block.location is not None

    def test_excel_with_bold_italic(self, tmp_path):
        """Process Excel file with bold and italic fonts"""
        file_path = tmp_path / "bold_italic.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = 'Bold text'
        ws['A1'].font = Font(bold=True)
        ws['A2'] = 'Italic text'
        ws['A2'].font = Font(italic=True)
        ws['A3'] = 'Bold italic'
        ws['A3'].font = Font(bold=True, italic=True)
        wb.save(file_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        assert len(blocks) == 3


# =============================================================================
# Tests: Skip Pattern Validation
# =============================================================================

class TestSkipPatterns:
    """Tests for text that should be skipped"""

    def test_excel_skips_numbers(self, tmp_path):
        """Numbers-only cells should be skipped"""
        file_path = tmp_path / "numbers.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = 12345
        ws['A2'] = 123.45
        ws['A3'] = -123
        ws['A4'] = 'Text content'
        wb.save(file_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        # Only 'Text content' should be extracted
        texts = [b.text for b in blocks]
        assert 'Text content' in texts
        assert len(blocks) == 1

    def test_excel_skips_emails(self, tmp_path):
        """Email addresses should be skipped"""
        file_path = tmp_path / "emails.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = 'test@example.com'
        ws['A2'] = 'user.name@domain.co.jp'
        ws['A3'] = 'Normal text'
        wb.save(file_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        texts = [b.text for b in blocks]
        assert 'Normal text' in texts

    def test_excel_skips_urls(self, tmp_path):
        """URLs should be skipped"""
        file_path = tmp_path / "urls.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = 'https://example.com'
        ws['A2'] = 'http://test.co.jp/path'
        ws['A3'] = 'www.example.com'
        ws['A4'] = 'Normal text'
        wb.save(file_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        texts = [b.text for b in blocks]
        assert 'Normal text' in texts


# =============================================================================
# Tests: File Info Extraction
# =============================================================================

class TestFileInfo:
    """Tests for file info extraction"""

    def test_excel_file_info(self, tmp_path):
        """ExcelProcessor returns correct file info"""
        file_path = tmp_path / "info_test.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = 'Test 1'
        ws['A2'] = 'Test 2'
        ws['A3'] = 'Test 3'
        wb.save(file_path)

        processor = ExcelProcessor()
        info = processor.get_file_info(file_path)

        assert info.path == file_path
        assert info.file_type == FileType.EXCEL
        assert info.size_bytes > 0

    def test_word_file_info(self, tmp_path):
        """WordProcessor returns correct file info"""
        file_path = tmp_path / "info_test.docx"
        doc = Document()
        doc.add_paragraph('Test 1')
        doc.add_paragraph('Test 2')
        doc.save(file_path)

        processor = WordProcessor()
        info = processor.get_file_info(file_path)

        assert info.path == file_path
        assert info.file_type == FileType.WORD
        assert info.size_bytes > 0


# =============================================================================
# Tests: Apply Translations
# =============================================================================

class TestApplyTranslations:
    """Tests for applying translations to files"""

    def test_excel_apply_translations(self, tmp_path):
        """Apply translations to Excel file"""
        input_path = tmp_path / "input.xlsx"
        output_path = tmp_path / "output.xlsx"

        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = 'Hello'
        ws['A2'] = 'World'
        wb.save(input_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(input_path))

        # Create translations
        translations = {
            blocks[0].id: 'こんにちは',
            blocks[1].id: '世界',
        }

        processor.apply_translations(input_path, output_path, translations)

        # Verify output
        assert output_path.exists()
        wb_out = openpyxl.load_workbook(output_path)
        ws_out = wb_out.active
        assert ws_out['A1'].value == 'こんにちは'
        assert ws_out['A2'].value == '世界'

    def test_word_apply_translations(self, tmp_path):
        """Apply translations to Word file"""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        doc = Document()
        doc.add_paragraph('Hello')
        doc.add_paragraph('World')
        doc.save(input_path)

        processor = WordProcessor()
        blocks = list(processor.extract_text_blocks(input_path))

        # Create translations
        translations = {
            blocks[0].id: 'こんにちは',
            blocks[1].id: '世界',
        }

        processor.apply_translations(input_path, output_path, translations)

        # Verify output
        assert output_path.exists()
        doc_out = Document(output_path)
        paragraphs = [p.text for p in doc_out.paragraphs if p.text.strip()]
        assert 'こんにちは' in paragraphs
        assert '世界' in paragraphs

    def test_apply_partial_translations(self, tmp_path):
        """Apply translations when some blocks are missing"""
        input_path = tmp_path / "input.xlsx"
        output_path = tmp_path / "output.xlsx"

        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = 'Hello'
        ws['A2'] = 'World'
        ws['A3'] = 'Test'
        wb.save(input_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(input_path))

        # Only translate first block
        translations = {
            blocks[0].id: 'こんにちは',
        }

        processor.apply_translations(input_path, output_path, translations)

        # Verify output - untranslated cells keep original
        wb_out = openpyxl.load_workbook(output_path)
        ws_out = wb_out.active
        assert ws_out['A1'].value == 'こんにちは'
        assert ws_out['A2'].value == 'World'  # Not translated
        assert ws_out['A3'].value == 'Test'   # Not translated


# =============================================================================
# Tests: Edge Cases in Location
# =============================================================================

class TestLocationHandling:
    """Tests for text block location handling"""

    def test_excel_location_format(self, tmp_path):
        """Excel locations include sheet and cell reference"""
        file_path = tmp_path / "location.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = 'Sheet1'
        ws['C5'] = 'Test'
        wb.save(file_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        assert len(blocks) == 1
        # Location should reference the cell
        assert 'C5' in blocks[0].location or '5' in blocks[0].location

    def test_excel_multisheet_locations(self, tmp_path):
        """Excel locations are unique across sheets"""
        file_path = tmp_path / "multi_location.xlsx"
        wb = openpyxl.Workbook()
        ws1 = wb.active
        ws1.title = 'First'
        ws1['A1'] = 'Text 1'
        ws2 = wb.create_sheet('Second')
        ws2['A1'] = 'Text 2'
        wb.save(file_path)

        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(file_path))

        assert len(blocks) == 2
        # Locations should be different
        assert blocks[0].location != blocks[1].location


# =============================================================================
# Tests: File Corruption Handling
# =============================================================================

class TestCorruptedFiles:
    """Tests for handling corrupted or invalid files"""

    def test_truncated_xlsx(self, tmp_path):
        """Handle truncated Excel file"""
        file_path = tmp_path / "truncated.xlsx"
        # Write partial file content
        file_path.write_bytes(b"PK\x03\x04truncated")

        processor = ExcelProcessor()

        with pytest.raises(Exception):
            processor.get_file_info(file_path)

    def test_wrong_extension(self, tmp_path):
        """Handle file with wrong extension"""
        # Create a text file but name it .xlsx
        file_path = tmp_path / "fake.xlsx"
        file_path.write_text("This is not an Excel file")

        processor = ExcelProcessor()

        with pytest.raises(Exception):
            processor.get_file_info(file_path)

    def test_nonexistent_file(self, tmp_path):
        """Handle non-existent file"""
        file_path = tmp_path / "does_not_exist.xlsx"

        processor = ExcelProcessor()

        with pytest.raises(Exception):
            processor.get_file_info(file_path)
