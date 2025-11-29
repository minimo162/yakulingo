# ecm_translate/processors/word_processor.py
"""
Processor for Word files (.docx, .doc).
"""

from pathlib import Path
from typing import Iterator
from docx import Document
from docx.shared import Pt

from .base import FileProcessor
from .translators import CellTranslator, ParagraphTranslator
from .font_manager import FontManager, FontTypeDetector
from ecm_translate.models.types import TextBlock, FileInfo, FileType


class WordProcessor(FileProcessor):
    """
    Processor for Word files (.docx, .doc).

    Translation targets:
    - Body paragraphs (ParagraphTranslator)
    - Table cells (CellTranslator - Excel-compatible)
    - Text boxes (ParagraphTranslator)

    Preserved:
    - Styles (headings, fonts)
    - Images and positions
    - Page layout
    - Lists (bullets, numbers)
    - Table formatting

    Note: Headers/Footers are NOT translated (excluded from processing)
    """

    def __init__(self):
        self.cell_translator = CellTranslator()
        self.para_translator = ParagraphTranslator()
        self.font_type_detector = FontTypeDetector()

    @property
    def file_type(self) -> FileType:
        return FileType.WORD

    @property
    def supported_extensions(self) -> list[str]:
        return ['.docx', '.doc']

    def get_file_info(self, file_path: Path) -> FileInfo:
        """Get Word file info"""
        doc = Document(file_path)

        text_count = 0

        # Count paragraphs
        for para in doc.paragraphs:
            if para.text and self.para_translator.should_translate(para.text):
                text_count += 1

        # Count table cells (Excel-compatible logic)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and self.cell_translator.should_translate(cell.text):
                        text_count += 1

        # Note: Headers/Footers are excluded from translation

        return FileInfo(
            path=file_path,
            file_type=FileType.WORD,
            size_bytes=file_path.stat().st_size,
            page_count=None,  # Requires full rendering
            text_block_count=text_count,
        )

    def extract_text_blocks(self, file_path: Path) -> Iterator[TextBlock]:
        """Extract text from paragraphs, tables"""
        doc = Document(file_path)

        # === Body Paragraphs ===
        for idx, para in enumerate(doc.paragraphs):
            if para.text and self.para_translator.should_translate(para.text):
                # Get font info from first run
                font_name = None
                font_size = 11.0
                if para.runs:
                    first_run = para.runs[0]
                    if first_run.font.name:
                        font_name = first_run.font.name
                    if first_run.font.size:
                        font_size = first_run.font.size.pt

                # Get dominant font if multiple runs
                if len(para.runs) > 1:
                    font_names = [r.font.name for r in para.runs if r.font.name]
                    if font_names:
                        font_name = self.font_type_detector.get_dominant_font(font_names)

                yield TextBlock(
                    id=f"para_{idx}",
                    text=para.text,
                    location=f"Paragraph {idx + 1}",
                    metadata={
                        'type': 'paragraph',
                        'index': idx,
                        'style': para.style.name if para.style else None,
                        'font_name': font_name,
                        'font_size': font_size,
                    }
                )

        # === Tables (Excel-compatible) ===
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text
                    if cell_text and self.cell_translator.should_translate(cell_text):
                        # Get font info from first paragraph's first run
                        font_name = None
                        font_size = 11.0
                        if cell.paragraphs and cell.paragraphs[0].runs:
                            first_run = cell.paragraphs[0].runs[0]
                            if first_run.font.name:
                                font_name = first_run.font.name
                            if first_run.font.size:
                                font_size = first_run.font.size.pt

                        yield TextBlock(
                            id=f"table_{table_idx}_r{row_idx}_c{cell_idx}",
                            text=cell_text,
                            location=f"Table {table_idx + 1}, Row {row_idx + 1}, Cell {cell_idx + 1}",
                            metadata={
                                'type': 'table_cell',
                                'table': table_idx,
                                'row': row_idx,
                                'col': cell_idx,
                                'font_name': font_name,
                                'font_size': font_size,
                            }
                        )

        # Note: Headers/Footers are excluded from translation

    def apply_translations(
        self,
        input_path: Path,
        output_path: Path,
        translations: dict[str, str],
        direction: str = "jp_to_en",
    ) -> None:
        """Apply translations while preserving formatting"""
        doc = Document(input_path)
        font_manager = FontManager(direction)

        # === Apply to paragraphs ===
        for idx, para in enumerate(doc.paragraphs):
            block_id = f"para_{idx}"
            if block_id in translations:
                self._apply_to_paragraph(para, translations[block_id], font_manager)

        # === Apply to tables ===
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    block_id = f"table_{table_idx}_r{row_idx}_c{cell_idx}"
                    if block_id in translations:
                        # Apply to first paragraph of cell
                        if cell.paragraphs:
                            self._apply_to_paragraph(
                                cell.paragraphs[0],
                                translations[block_id],
                                font_manager
                            )
                            # Clear remaining paragraphs if any
                            for para in cell.paragraphs[1:]:
                                for run in para.runs:
                                    run.text = ""

        # Note: Headers/Footers are excluded from translation

        doc.save(output_path)

    def _apply_to_paragraph(self, para, translated_text: str, font_manager: FontManager) -> None:
        """
        Apply translation to paragraph, preserving paragraph style.

        Strategy:
        - Keep first run's formatting
        - Set translated text to first run
        - Clear remaining runs
        - Apply new font based on direction
        """
        if para.runs:
            first_run = para.runs[0]

            # Get original font info
            original_font_name = first_run.font.name
            original_font_size = first_run.font.size.pt if first_run.font.size else 11.0

            # Get new font settings
            new_font_name, new_font_size = font_manager.select_font(
                original_font_name,
                original_font_size
            )

            # Apply translation
            first_run.text = translated_text

            # Apply new font
            first_run.font.name = new_font_name
            first_run.font.size = Pt(new_font_size)

            # Clear remaining runs
            for run in para.runs[1:]:
                run.text = ""
        else:
            para.text = translated_text
