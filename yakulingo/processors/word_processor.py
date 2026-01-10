# yakulingo/processors/word_processor.py
"""
Processor for Word files (.docx).
"""

import logging
import re
import shutil
import sys
import tempfile
import threading
import zipfile
from contextlib import contextmanager
from functools import lru_cache
from pathlib import Path
from typing import Iterator, Optional
from xml.etree import ElementTree as ET
from docx import Document
from docx.shared import Pt

from .translators import CellTranslator, ParagraphTranslator
from .font_manager import FontManager
from yakulingo.models.types import TextBlock, FileInfo, FileType, SectionDetail
from .base import FileProcessor

# Module logger
logger = logging.getLogger(__name__)


# =============================================================================
# TextBox Extraction via XML (python-docx doesn't support this)
# =============================================================================
# XML namespaces used in Word documents
WORD_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "v": "urn:schemas-microsoft-com:vml",
}

# =============================================================================
# Page mapping via Word COM (for page-select translation)
# =============================================================================
# NOTE: .docx has no fixed "pages" without layout rendering.
# We use Microsoft Word COM automation on Windows to obtain true page numbers.
_pythoncom = None

# Word constants (avoid win32com.client.constants to keep things robust without gencache)
_WD_MAIN_TEXT_STORY = 1
_WD_STATISTIC_PAGES = 2
_WD_ACTIVE_END_PAGE_NUMBER = 3
_WD_WITHIN_TABLE = 12

# Office constants
_MSO_AUTOMATION_SECURITY_FORCE_DISABLE = 3


def _get_pythoncom():
    """Lazy import pythoncom (Windows COM library)."""
    global _pythoncom
    if _pythoncom is None and sys.platform == "win32":
        try:
            import pythoncom

            _pythoncom = pythoncom
        except ImportError:
            logger.debug("pythoncom not available")
    return _pythoncom


@contextmanager
def _com_initialized():
    """
    Ensure COM is initialized for the current thread (required in worker threads).
    On non-Windows platforms this is a no-op.
    """
    pythoncom = _get_pythoncom()
    initialized = False

    if pythoncom is not None:
        try:
            thread_id = threading.current_thread().ident
            try:
                hr = pythoncom.CoInitializeEx(pythoncom.COINIT_APARTMENTTHREADED)
                # S_OK (0)=newly initialized, S_FALSE (1)=already initialized
                initialized = hr == 0 or hr is None
                logger.debug(
                    "Word COM initialized (STA) thread=%s hr=%s will_uninit=%s",
                    thread_id,
                    hr,
                    initialized,
                )
            except Exception:
                hr = pythoncom.CoInitialize()
                initialized = hr == 0 or hr is None
                logger.debug(
                    "Word COM initialized (fallback) thread=%s hr=%s will_uninit=%s",
                    thread_id,
                    hr,
                    initialized,
                )
        except Exception as e:
            logger.debug("Word COM initialization skipped: %s", e)

    try:
        yield
    finally:
        if initialized and pythoncom is not None:
            try:
                thread_id = threading.current_thread().ident
                pythoncom.CoUninitialize()
                logger.debug("Word COM uninitialized thread=%s", thread_id)
            except Exception as e:
                logger.debug("Word COM uninitialization failed: %s", e)


def _try_get_word_page_data_via_com(
    file_path: Path,
) -> tuple[Optional[int], dict[str, int]]:
    """
    Get page count and block_id -> page_idx mapping via Microsoft Word COM.

    Returns:
        (page_count, page_map) where page_idx is 0-based.
        On failure: (None, {}).
    """
    if sys.platform != "win32":
        return None, {}

    try:
        import win32com.client  # type: ignore[import-not-found]
    except ImportError:
        logger.debug("win32com not available: Word page mapping disabled")
        return None, {}

    with _com_initialized():
        word = None
        doc = None
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            try:
                word.AutomationSecurity = _MSO_AUTOMATION_SECURITY_FORCE_DISABLE
            except Exception:
                pass

            doc = word.Documents.Open(
                str(file_path),
                ReadOnly=True,
                AddToRecentFiles=False,
                Visible=False,
            )

            try:
                doc.Repaginate()
            except Exception:
                pass

            page_count: Optional[int]
            try:
                page_count = int(doc.ComputeStatistics(_WD_STATISTIC_PAGES))
            except Exception:
                page_count = None

            try:
                main_range = doc.StoryRanges(_WD_MAIN_TEXT_STORY)
            except Exception:
                main_range = doc.Range()

            page_map: dict[str, int] = {}

            # --- Body paragraphs (exclude table paragraphs to match python-docx Document.paragraphs) ---
            para_idx = 0
            try:
                paragraphs = main_range.Paragraphs
                for i in range(1, int(paragraphs.Count) + 1):
                    para = paragraphs(i)
                    try:
                        if bool(para.Range.Information(_WD_WITHIN_TABLE)):
                            continue
                    except Exception:
                        pass

                    try:
                        page_num = int(
                            para.Range.Information(_WD_ACTIVE_END_PAGE_NUMBER)
                        )
                    except Exception:
                        page_num = 1
                    page_map[f"para_{para_idx}"] = max(0, page_num - 1)
                    para_idx += 1
            except Exception as e:
                logger.debug("Failed to read Word paragraph pages via COM: %s", e)

            # --- Tables (top-level only, to match python-docx Document.tables) ---
            table_idx = 0
            try:
                tables = main_range.Tables
                for i in range(1, int(tables.Count) + 1):
                    table = tables(i)
                    try:
                        nesting_level = int(getattr(table, "NestingLevel", 1))
                    except Exception:
                        nesting_level = 1
                    if nesting_level != 1:
                        continue

                    try:
                        row_count = int(table.Rows.Count)
                        col_count = int(table.Columns.Count)
                    except Exception:
                        row_count = 0
                        col_count = 0

                    for row in range(1, row_count + 1):
                        for col in range(1, col_count + 1):
                            try:
                                cell = table.Cell(row, col)
                            except Exception:
                                continue

                            try:
                                page_num = int(
                                    cell.Range.Information(_WD_ACTIVE_END_PAGE_NUMBER)
                                )
                            except Exception:
                                page_num = 1

                            page_map[f"table_{table_idx}_r{row - 1}_c{col - 1}"] = max(
                                0, page_num - 1
                            )

                    table_idx += 1
            except Exception as e:
                logger.debug("Failed to read Word table pages via COM: %s", e)

            # --- Text boxes / Shapes (best effort) ---
            textbox_idx = 0
            try:
                shapes_with_text = []
                shapes = doc.Shapes
                for i in range(1, int(shapes.Count) + 1):
                    shape = shapes(i)
                    try:
                        anchor = shape.Anchor
                        if int(anchor.StoryType) != _WD_MAIN_TEXT_STORY:
                            continue
                    except Exception:
                        anchor = None

                    try:
                        if not bool(shape.TextFrame.HasText):
                            continue
                        text = str(shape.TextFrame.TextRange.Text).strip()
                    except Exception:
                        continue

                    if not text:
                        continue

                    try:
                        page_num = int(
                            (anchor or shape.Anchor).Information(
                                _WD_ACTIVE_END_PAGE_NUMBER
                            )
                        )
                    except Exception:
                        page_num = 1

                    try:
                        anchor_start = int((anchor or shape.Anchor).Start)
                    except Exception:
                        anchor_start = 0

                    shapes_with_text.append((anchor_start, page_num))

                shapes_with_text.sort(key=lambda x: x[0])

                for _, page_num in shapes_with_text:
                    page_map[f"textbox_{textbox_idx}"] = max(0, int(page_num) - 1)
                    textbox_idx += 1
            except Exception as e:
                logger.debug("Failed to read Word shape pages via COM: %s", e)

            return page_count, page_map

        except Exception as e:
            logger.debug("Word COM page mapping failed for '%s': %s", file_path, e)
            return None, {}

        finally:
            try:
                if doc is not None:
                    doc.Close(False)
            except Exception:
                pass
            try:
                if word is not None:
                    # Do not save changes (we opened ReadOnly=True anyway).
                    word.Quit(False)
            except Exception:
                pass
            # Help GC release COM proxies
            try:
                del doc
            except Exception:
                pass
            try:
                del word
            except Exception:
                pass


@lru_cache(maxsize=8)
def _get_word_page_data_cached(
    path_str: str, mtime_ns: int, size_bytes: int
) -> tuple[Optional[int], dict[str, int]]:
    # mtime_ns/size_bytes are included to invalidate cache on file changes.
    return _try_get_word_page_data_via_com(Path(path_str))


def _extract_text_from_txbx_content(txbx_content) -> str:
    """
    Extract text from a txbxContent XML element.

    This helper function extracts text from both modern (wps:txbx) and
    legacy (v:textbox) Word textbox elements.

    Args:
        txbx_content: w:txbxContent XML element

    Returns:
        Extracted text joined by newlines, or empty string if no text
    """
    if txbx_content is None:
        return ""

    text_parts = []
    for p in txbx_content.findall(".//w:p", WORD_NS):
        para_text = []
        for t in p.findall(".//w:t", WORD_NS):
            if t.text:
                para_text.append(t.text)
        if para_text:
            text_parts.append("".join(para_text))

    if text_parts:
        return "\n".join(text_parts).strip()
    return ""


def _extract_textboxes_from_docx(file_path: Path) -> list[dict]:
    """
    Extract TextBox content from docx file by parsing XML directly.

    python-docx doesn't support TextBox text extraction, so we parse
    the word/document.xml file directly from the docx archive.

    Args:
        file_path: Path to docx file

    Returns:
        List of dicts with 'textbox_index', 'text' keys
    """
    textboxes = []

    try:
        with zipfile.ZipFile(file_path, "r") as zf:
            # Read main document
            if "word/document.xml" not in zf.namelist():
                return textboxes

            xml_content = zf.read("word/document.xml")
            root = ET.fromstring(xml_content)

            textbox_index = 0

            # Method 1: Modern Word textboxes (wps:txbx)
            for txbx in root.findall(".//wps:txbx", WORD_NS):
                txbx_content = txbx.find(".//w:txbxContent", WORD_NS)
                full_text = _extract_text_from_txbx_content(txbx_content)
                if full_text:
                    textboxes.append(
                        {
                            "textbox_index": textbox_index,
                            "text": full_text,
                        }
                    )
                    textbox_index += 1

            # Method 2: Legacy VML textboxes (v:textbox)
            for textbox in root.findall(".//v:textbox", WORD_NS):
                txbx_content = textbox.find(".//w:txbxContent", WORD_NS)
                full_text = _extract_text_from_txbx_content(txbx_content)
                if full_text:
                    textboxes.append(
                        {
                            "textbox_index": textbox_index,
                            "text": full_text,
                        }
                    )
                    textbox_index += 1

    except (zipfile.BadZipFile, ET.ParseError):
        pass

    return textboxes


def _apply_textbox_translations_to_docx(
    output_path: Path,
    translations: dict[str, str],
) -> None:
    """
    Apply translations to TextBox content by modifying XML directly.

    Args:
        output_path: Output docx file (must already exist from python-docx save)
        translations: Dict mapping textbox IDs to translated text
    """
    # Filter textbox translations
    textbox_translations = {
        k: v for k, v in translations.items() if k.startswith("textbox_")
    }

    if not textbox_translations:
        return

    # Parse textbox IDs
    textbox_map = {}  # textbox_index -> translated_text
    for block_id, translated in textbox_translations.items():
        match = re.match(r"textbox_(\d+)", block_id)
        if match:
            tb_idx = int(match.group(1))
            textbox_map[tb_idx] = translated

    if not textbox_map:
        return

    # Create a temporary copy to work with
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_docx = Path(temp_dir) / "temp.docx"
        shutil.copy(output_path, temp_docx)

        try:
            with zipfile.ZipFile(temp_docx, "r") as zf_in:
                with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf_out:
                    for item in zf_in.namelist():
                        content = zf_in.read(item)

                        if item == "word/document.xml":
                            content = _modify_docx_textboxes(content, textbox_map)

                        zf_out.writestr(item, content)

        except (zipfile.BadZipFile, ET.ParseError, OSError) as e:
            logger.warning("Error modifying docx textboxes: %s", e)
            shutil.copy(temp_docx, output_path)


def _modify_docx_textboxes(xml_content: bytes, translations: dict[int, str]) -> bytes:
    """
    Modify document XML to apply textbox translations.
    """
    try:
        root = ET.fromstring(xml_content)
        textbox_index = 0

        # Process modern textboxes (wps:txbx)
        for txbx in root.findall(".//wps:txbx", WORD_NS):
            txbxContent = txbx.find(".//w:txbxContent", WORD_NS)
            if txbxContent is not None:
                # Check if has text
                has_text = any(
                    t.text
                    for p in txbxContent.findall(".//w:p", WORD_NS)
                    for t in p.findall(".//w:t", WORD_NS)
                    if t.text
                )

                if has_text and textbox_index in translations:
                    # Apply translation to first text element, clear others
                    first_t = None
                    for p in txbxContent.findall(".//w:p", WORD_NS):
                        for t in p.findall(".//w:t", WORD_NS):
                            if first_t is None:
                                first_t = t
                                t.text = translations[textbox_index]
                            else:
                                t.text = ""

                if has_text:
                    textbox_index += 1

        # Process legacy VML textboxes
        for textbox in root.findall(".//v:textbox", WORD_NS):
            txbxContent = textbox.find(".//w:txbxContent", WORD_NS)
            if txbxContent is not None:
                has_text = any(
                    t.text
                    for p in txbxContent.findall(".//w:p", WORD_NS)
                    for t in p.findall(".//w:t", WORD_NS)
                    if t.text
                )

                if has_text and textbox_index in translations:
                    first_t = None
                    for p in txbxContent.findall(".//w:p", WORD_NS):
                        for t in p.findall(".//w:t", WORD_NS):
                            if first_t is None:
                                first_t = t
                                t.text = translations[textbox_index]
                            else:
                                t.text = ""

                if has_text:
                    textbox_index += 1

        return ET.tostring(root, encoding="unicode").encode("utf-8")

    except ET.ParseError:
        return xml_content


class WordProcessor(FileProcessor):
    """
    Processor for Word files (.docx, .doc).

    Translation targets:
    - Body paragraphs (ParagraphTranslator)
    - Table cells (CellTranslator - Excel-compatible)
    - Text boxes (ParagraphTranslator)

    Preserved:
    - Styles (headings, fonts)
    - Images and positions
    - Page layout
    - Lists (bullets, numbers)
    - Table formatting

    Note: Headers/Footers are NOT translated (excluded from processing)
    """

    def __init__(self):
        self.cell_translator = CellTranslator()
        self.para_translator = ParagraphTranslator()

    @property
    def file_type(self) -> FileType:
        return FileType.WORD

    @property
    def supported_extensions(self) -> list[str]:
        # Note: .doc (legacy format) is not supported by python-docx
        # Only .docx (Office Open XML) is supported
        return [".docx"]

    def get_file_info(self, file_path: Path) -> FileInfo:
        """Get Word file info (page count via Word COM when available)."""
        page_count = None
        page_map: dict[str, int] = {}

        try:
            stat = file_path.stat()
            page_count, page_map = _get_word_page_data_cached(
                str(file_path), stat.st_mtime_ns, stat.st_size
            )
        except Exception:
            page_count = None
            page_map = {}

        if page_count is None and page_map:
            try:
                page_count = max(page_map.values()) + 1
            except ValueError:
                page_count = None

        if page_count and page_count > 1:
            section_details = [
                SectionDetail(index=idx, name=f"ページ {idx + 1}")
                for idx in range(page_count)
            ]
        else:
            section_details = [SectionDetail(index=0, name="全体")]

        return FileInfo(
            path=file_path,
            file_type=FileType.WORD,
            size_bytes=file_path.stat().st_size,
            page_count=page_count,
            section_details=section_details,
        )

    def extract_sample_text_fast(
        self, file_path: Path, max_chars: int = 500
    ) -> Optional[str]:
        """Extract a text sample for language detection without full document load.

        This method uses iterparse to stream document.xml directly from the docx
        archive, avoiding the overhead of loading the entire XML into memory.
        Large Word files can have document.xml of tens of MB, so streaming
        is essential for fast language detection.

        Args:
            file_path: Path to the Word file
            max_chars: Maximum characters to extract (default 500)

        Returns:
            Sample text string or None if extraction fails
        """
        if file_path.suffix.lower() != ".docx":
            return None

        try:
            texts = []
            total_chars = 0

            with zipfile.ZipFile(file_path, "r") as zf:
                if "word/document.xml" not in zf.namelist():
                    logger.debug("No document.xml found in docx")
                    return None

                # Use iterparse for streaming XML parsing (avoids loading entire XML into memory)
                # This is critical for large Word files where document.xml can be huge
                with zf.open("word/document.xml") as xml_file:
                    # Word namespace
                    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                    t_tag = f"{{{ns}}}t"

                    for event, elem in ET.iterparse(xml_file, events=("end",)):
                        if elem.tag == t_tag and elem.text:
                            text = elem.text.strip()
                            if text and len(text) > 1:  # Skip single chars
                                texts.append(text)
                                total_chars += len(text)
                                if total_chars >= max_chars:
                                    break
                        # Clear element to free memory
                        elem.clear()

            if texts:
                result = " ".join(texts)[:max_chars]
                logger.debug(
                    "Word fast sample extraction: %d chars from %d text runs",
                    len(result),
                    len(texts),
                )
                return result

            return None

        except (zipfile.BadZipFile, ET.ParseError, KeyError) as e:
            logger.debug("Word fast sample extraction failed: %s", e)
            return None

    def extract_text_blocks(
        self, file_path: Path, output_language: str = "en"
    ) -> Iterator[TextBlock]:
        """Extract text from paragraphs, tables, and textboxes (.docx only)

        Args:
            file_path: Path to the Word file
            output_language: "en" for JP→EN, "jp" for EN→JP translation
        """
        page_map: dict[str, int] = {}
        try:
            stat = file_path.stat()
            _, page_map = _get_word_page_data_cached(
                str(file_path), stat.st_mtime_ns, stat.st_size
            )
        except Exception:
            page_map = {}

        doc = Document(file_path)
        try:
            # === Body Paragraphs ===
            for idx, para in enumerate(doc.paragraphs):
                if para.text and self.para_translator.should_translate(
                    para.text, output_language
                ):
                    # Get font info from first run
                    font_name = None
                    font_size = 11.0
                    if para.runs:
                        first_run = para.runs[0]
                        if first_run.font.name:
                            font_name = first_run.font.name
                        if first_run.font.size:
                            font_size = first_run.font.size.pt

                    # Get font from first valid run if multiple runs
                    if len(para.runs) > 1:
                        font_names = [r.font.name for r in para.runs if r.font.name]
                        if font_names:
                            font_name = font_names[0]

                    block_id = f"para_{idx}"
                    page_idx = page_map.get(block_id)

                    yield TextBlock(
                        id=block_id,
                        text=para.text,
                        location=f"Paragraph {idx + 1}",
                        metadata={
                            "type": "paragraph",
                            "index": idx,
                            "style": para.style.name if para.style else None,
                            "font_name": font_name,
                            "font_size": font_size,
                            **({"page_idx": page_idx} if page_idx is not None else {}),
                        },
                    )

            # === Tables (Excel-compatible) ===
            # Track processed cells to avoid extracting merged cells multiple times
            for table_idx, table in enumerate(doc.tables):
                processed_cells = set()
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        # Use the underlying tc element to deduplicate merged cells
                        cell_key = cell._tc
                        if cell_key in processed_cells:
                            continue
                        processed_cells.add(cell_key)

                        cell_text = cell.text
                        if cell_text and self.cell_translator.should_translate(
                            cell_text, output_language
                        ):
                            # Get font info from first paragraph's first run
                            font_name = None
                            font_size = 11.0
                            if cell.paragraphs and cell.paragraphs[0].runs:
                                first_run = cell.paragraphs[0].runs[0]
                                if first_run.font.name:
                                    font_name = first_run.font.name
                                if first_run.font.size:
                                    font_size = first_run.font.size.pt

                            block_id = f"table_{table_idx}_r{row_idx}_c{cell_idx}"
                            page_idx = page_map.get(block_id)

                            yield TextBlock(
                                id=block_id,
                                text=cell_text,
                                location=f"Table {table_idx + 1}, Row {row_idx + 1}, Cell {cell_idx + 1}",
                                metadata={
                                    "type": "table_cell",
                                    "table": table_idx,
                                    "row": row_idx,
                                    "col": cell_idx,
                                    "font_name": font_name,
                                    "font_size": font_size,
                                    **(
                                        {"page_idx": page_idx}
                                        if page_idx is not None
                                        else {}
                                    ),
                                },
                            )

            # Note: Headers/Footers are excluded from translation

            # === TextBoxes (via XML parsing, docx only) ===
            if str(file_path).lower().endswith(".docx"):
                textboxes = _extract_textboxes_from_docx(file_path)
                for tb in textboxes:
                    text = tb["text"]
                    tb_idx = tb["textbox_index"]

                    if self.para_translator.should_translate(text, output_language):
                        block_id = f"textbox_{tb_idx}"
                        page_idx = page_map.get(block_id)
                        yield TextBlock(
                            id=block_id,
                            text=text,
                            location=f"TextBox {tb_idx + 1}",
                            metadata={
                                "type": "textbox",
                                "textbox": tb_idx,
                                **(
                                    {"page_idx": page_idx}
                                    if page_idx is not None
                                    else {}
                                ),
                            },
                        )
        finally:
            # python-docx doesn't have close(), but we can help GC by deleting the reference
            del doc

    def apply_translations(
        self,
        input_path: Path,
        output_path: Path,
        translations: dict[str, str],
        direction: str = "jp_to_en",
        settings=None,
        selected_sections: Optional[list[int]] = None,
        text_blocks=None,  # Accept but unused (for API compatibility)
    ) -> None:
        """Apply translations while preserving formatting.

        Note: selected_sections and text_blocks are accepted for API consistency
        but not used for Word documents (Word doesn't have the same positioning issues).
        """
        doc = Document(input_path)
        try:
            font_manager = FontManager(direction, settings)

            # === Apply to paragraphs ===
            for idx, para in enumerate(doc.paragraphs):
                block_id = f"para_{idx}"
                if block_id in translations:
                    self._apply_to_paragraph(para, translations[block_id], font_manager)

            # === Apply to tables ===
            # Track processed cells to avoid applying to merged cells multiple times
            for table_idx, table in enumerate(doc.tables):
                processed_cells = set()
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_key = cell._tc
                        if cell_key in processed_cells:
                            continue
                        processed_cells.add(cell_key)

                        block_id = f"table_{table_idx}_r{row_idx}_c{cell_idx}"
                        if block_id in translations:
                            # Apply to first paragraph of cell
                            if cell.paragraphs:
                                self._apply_to_paragraph(
                                    cell.paragraphs[0],
                                    translations[block_id],
                                    font_manager,
                                )
                                # Clear remaining paragraphs if any
                                for para in cell.paragraphs[1:]:
                                    for run in para.runs:
                                        run.text = ""

            # Note: Headers/Footers are excluded from translation

            doc.save(output_path)
        finally:
            # python-docx doesn't have close(), but we can help GC by deleting the reference
            del doc

        # Apply TextBox translations via XML manipulation (python-docx doesn't support this)
        # Only for .docx files
        if str(input_path).lower().endswith(".docx"):
            _apply_textbox_translations_to_docx(output_path, translations)

    def _apply_to_paragraph(
        self, para, translated_text: str, font_manager: FontManager
    ) -> None:
        """
        Apply translation to paragraph, preserving paragraph style.

        Strategy:
        - Keep first run's formatting
        - Set translated text to first run
        - Clear remaining runs
        - Apply new font based on direction
        """
        if para.runs:
            first_run = para.runs[0]

            # Get original font info
            original_font_name = first_run.font.name
            original_font_size = first_run.font.size.pt if first_run.font.size else 11.0

            # Get new font settings
            new_font_name, new_font_size = font_manager.select_font(
                original_font_name, original_font_size
            )

            # Apply translation
            first_run.text = translated_text

            # Apply new font
            first_run.font.name = new_font_name
            first_run.font.size = Pt(new_font_size)

            # Clear remaining runs
            for run in para.runs[1:]:
                run.text = ""
        else:
            # No runs - add text via a new run (para.text is read-only)
            para.add_run().text = translated_text

    def create_bilingual_document(
        self,
        original_path: Path,
        translated_path: Path,
        output_path: Path,
    ) -> dict[str, int]:
        """
        Create a bilingual document with original and translated content.

        Output format:
            Original content
            --- Page Break ---
            Translated content

        Args:
            original_path: Path to the original document
            translated_path: Path to the translated document
            output_path: Path to save the bilingual document

        Returns:
            dict with original_paragraphs, translated_paragraphs counts
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        original_doc = Document(original_path)
        translated_doc = Document(translated_path)
        try:
            # Count paragraphs
            original_count = len(original_doc.paragraphs)
            translated_count = len(translated_doc.paragraphs)

            # Add separator heading before translated content
            separator_para = original_doc.add_paragraph()
            separator_para.add_run("─" * 40)
            separator_para.alignment = 1  # Center

            # Add page break
            page_break_para = original_doc.add_paragraph()
            run = page_break_para.add_run()
            run._r.append(OxmlElement("w:br", {qn("w:type"): "page"}))

            # Add translation header
            header_para = original_doc.add_paragraph()
            header_run = header_para.add_run("【翻訳】")
            header_run.bold = True

            # Copy translated paragraphs
            for para in translated_doc.paragraphs:
                new_para = original_doc.add_paragraph()
                # Copy style
                new_para.style = para.style
                new_para.alignment = para.alignment
                # Copy runs
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    if run.font.bold:
                        new_run.font.bold = run.font.bold
                    if run.font.italic:
                        new_run.font.italic = run.font.italic
                    if run.font.name:
                        new_run.font.name = run.font.name
                    if run.font.size:
                        new_run.font.size = run.font.size

            # Copy translated tables
            for table in translated_doc.tables:
                # Add table to original doc
                new_table = original_doc.add_table(
                    rows=len(table.rows), cols=len(table.columns)
                )
                new_table.style = table.style

                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        new_cell = new_table.cell(row_idx, col_idx)
                        # Copy cell content
                        for para_idx, para in enumerate(cell.paragraphs):
                            if para_idx == 0:
                                target_para = new_cell.paragraphs[0]
                            else:
                                target_para = new_cell.add_paragraph()
                            for run in para.runs:
                                target_para.add_run(run.text)

            original_doc.save(output_path)

            return {
                "original_paragraphs": original_count,
                "translated_paragraphs": translated_count,
                "total_paragraphs": original_count + translated_count,
            }
        finally:
            # python-docx doesn't have close(), but we can help GC by deleting the references
            del original_doc
            del translated_doc
