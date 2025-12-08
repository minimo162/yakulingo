# tests/test_word_processor.py
"""Tests for yakulingo.processors.word_processor"""

import pytest
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

from yakulingo.processors.word_processor import WORD_NS, WordProcessor
from yakulingo.models.types import FileType


# --- Fixtures ---

@pytest.fixture
def processor():
    """WordProcessor instance"""
    return WordProcessor()


@pytest.fixture
def sample_docx(tmp_path):
    """Create a sample Word document with paragraphs"""
    file_path = tmp_path / "sample.docx"
    doc = Document()

    # Add paragraphs
    doc.add_paragraph("これは最初の段落です。")
    doc.add_paragraph("This is the second paragraph.")
    doc.add_paragraph("12345")  # Numbers only - should skip
    doc.add_paragraph("test@example.com")  # Email - should skip

    doc.save(file_path)
    return file_path


@pytest.fixture
def docx_with_table(tmp_path):
    """Create Word document with a table"""
    file_path = tmp_path / "table.docx"
    doc = Document()

    doc.add_paragraph("Document with table")

    # Add table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "ヘッダー1"
    table.cell(0, 1).text = "Header 2"
    table.cell(1, 0).text = "12345"  # Should skip
    table.cell(1, 1).text = "データ"

    doc.save(file_path)
    return file_path


@pytest.fixture
def docx_with_font(tmp_path):
    """Create Word document with specific fonts"""
    file_path = tmp_path / "font.docx"
    doc = Document()

    # Paragraph with specific font
    para = doc.add_paragraph()
    run = para.add_run("フォント設定テスト")
    run.font.name = "MS Mincho"
    run.font.size = Pt(14)

    doc.save(file_path)
    return file_path


@pytest.fixture
def empty_docx(tmp_path):
    """Create empty Word document"""
    file_path = tmp_path / "empty.docx"
    doc = Document()
    doc.save(file_path)
    return file_path


# --- Tests: Properties ---

class TestWordProcessorProperties:
    """Test WordProcessor properties"""

    def test_file_type(self, processor):
        assert processor.file_type == FileType.WORD

    def test_supported_extensions(self, processor):
        extensions = processor.supported_extensions
        assert ".docx" in extensions
        # .doc (legacy format) is not supported by python-docx
        assert ".doc" not in extensions


# --- Tests: get_file_info ---

class TestWordProcessorGetFileInfo:
    """Test WordProcessor.get_file_info()"""

    def test_file_info_basic(self, processor, sample_docx):
        """Basic file info retrieval"""
        info = processor.get_file_info(sample_docx)

        assert info.path == sample_docx
        assert info.file_type == FileType.WORD
        assert info.size_bytes > 0

    def test_file_info_with_table(self, processor, docx_with_table):
        """File info basic attributes"""
        info = processor.get_file_info(docx_with_table)
        assert info.file_type == FileType.WORD

    def test_file_info_empty(self, processor, empty_docx):
        """File info for empty document"""
        info = processor.get_file_info(empty_docx)
        assert info.file_type == FileType.WORD


# --- Tests: extract_text_blocks ---

class TestWordProcessorExtractTextBlocks:
    """Test WordProcessor.extract_text_blocks()"""

    def test_extracts_paragraphs(self, processor, sample_docx):
        """Extracts translatable paragraphs (Japanese only)"""
        blocks = list(processor.extract_text_blocks(sample_docx))

        # Should have 1 block (only Japanese text, English-only is skipped)
        assert len(blocks) == 1

        texts = [b.text for b in blocks]
        assert "これは最初の段落です。" in texts
        # Note: "This is the second paragraph." is skipped as English-only

    def test_skips_numbers_and_emails(self, processor, sample_docx):
        """Skips non-translatable content"""
        blocks = list(processor.extract_text_blocks(sample_docx))
        texts = [b.text for b in blocks]

        assert "12345" not in texts
        assert "test@example.com" not in texts

    def test_extracts_table_cells(self, processor, docx_with_table):
        """Extracts table cells (Japanese only)"""
        blocks = list(processor.extract_text_blocks(docx_with_table))

        # Filter table cells
        table_blocks = [b for b in blocks if b.metadata.get("type") == "table_cell"]
        # Only Japanese cells are extracted
        assert len(table_blocks) == 2

        table_texts = [b.text for b in table_blocks]
        assert "ヘッダー1" in table_texts
        # Note: "Header 2" is skipped as English-only
        assert "データ" in table_texts
        assert "12345" not in table_texts  # Skipped

    def test_block_metadata_paragraph(self, processor, sample_docx):
        """Paragraph blocks have correct metadata"""
        blocks = list(processor.extract_text_blocks(sample_docx))

        para_block = blocks[0]
        assert para_block.metadata["type"] == "paragraph"
        assert "index" in para_block.metadata

    def test_block_metadata_table(self, processor, docx_with_table):
        """Table cell blocks have correct metadata"""
        blocks = list(processor.extract_text_blocks(docx_with_table))

        table_block = next(b for b in blocks if b.metadata.get("type") == "table_cell")
        assert "table" in table_block.metadata
        assert "row" in table_block.metadata
        assert "col" in table_block.metadata

    def test_block_ids_unique(self, processor, docx_with_table):
        """All block IDs are unique"""
        blocks = list(processor.extract_text_blocks(docx_with_table))
        ids = [b.id for b in blocks]
        assert len(ids) == len(set(ids))

    def test_extracts_font_info(self, processor, docx_with_font):
        """Extracts font information"""
        blocks = list(processor.extract_text_blocks(docx_with_font))

        assert len(blocks) == 1
        block = blocks[0]
        assert block.metadata.get("font_name") == "MS Mincho"
        assert block.metadata.get("font_size") == 14.0

    def test_empty_file_returns_no_blocks(self, processor, empty_docx):
        """Empty file yields no blocks"""
        blocks = list(processor.extract_text_blocks(empty_docx))
        assert blocks == []


class TestWordProcessorTextBoxes:
    """TextBox extraction and translation application"""

    def _add_textbox(self, doc_path: Path, text: str) -> None:
        """Inject a simple textbox paragraph into an existing docx file."""
        with zipfile.ZipFile(doc_path, "r") as zf:
            contents = {name: zf.read(name) for name in zf.namelist()}

        root = ET.fromstring(contents["word/document.xml"])

        ET.register_namespace("w", WORD_NS["w"])
        ET.register_namespace("wp", WORD_NS["wp"])
        ET.register_namespace("wps", WORD_NS["wps"])
        ET.register_namespace("a", WORD_NS["a"])
        ET.register_namespace("mc", WORD_NS["mc"])
        ET.register_namespace("v", WORD_NS["v"])

        body = root.find("w:body", WORD_NS)
        assert body is not None

        p = ET.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p")
        r = ET.SubElement(p, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
        drawing = ET.SubElement(r, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing")
        inline = ET.SubElement(drawing, "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline")
        graphic = ET.SubElement(inline, "{http://schemas.openxmlformats.org/drawingml/2006/main}graphic")
        graphicData = ET.SubElement(
            graphic,
            "{http://schemas.openxmlformats.org/drawingml/2006/main}graphicData",
            {
                "{http://schemas.openxmlformats.org/drawingml/2006/main}uri": (
                    "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
                )
            },
        )
        wsp = ET.SubElement(
            graphicData, "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}wsp"
        )
        txbx = ET.SubElement(
            wsp, "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx"
        )
        txbxContent = ET.SubElement(
            txbx, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent"
        )
        inner_p = ET.SubElement(
            txbxContent, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
        )
        inner_r = ET.SubElement(
            inner_p, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
        )
        t = ET.SubElement(inner_r, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
        t.text = text

        body.append(p)

        contents["word/document.xml"] = ET.tostring(root, encoding="utf-8")

        with zipfile.ZipFile(doc_path, "w", zipfile.ZIP_DEFLATED) as zf_out:
            for name, data in contents.items():
                zf_out.writestr(name, data)

    def test_extracts_textbox_content(self, processor, tmp_path):
        """TextBox content is included in extracted text blocks."""
        file_path = tmp_path / "textbox.docx"
        doc = Document()
        doc.add_paragraph("通常の段落")
        doc.save(file_path)

        self._add_textbox(file_path, "テキストボックスの内容")

        blocks = list(processor.extract_text_blocks(file_path))
        textbox_blocks = [b for b in blocks if b.metadata.get("type") == "textbox"]

        assert len(textbox_blocks) == 1
        assert textbox_blocks[0].text == "テキストボックスの内容"

    def test_applies_textbox_translation(self, processor, tmp_path):
        """Translations are applied to textbox content in the output docx."""
        input_path = tmp_path / "textbox.docx"
        doc = Document()
        doc.add_paragraph("通常の段落")
        doc.save(input_path)

        self._add_textbox(input_path, "未翻訳")

        translations = {
            "para_0": "Standard paragraph",
            "textbox_0": "Translated textbox",
        }

        output_path = tmp_path / "translated.docx"
        processor.apply_translations(
            input_path,
            output_path,
            translations,
            direction="jp_to_en",
        )

        with zipfile.ZipFile(output_path, "r") as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")

        assert "Translated textbox" in document_xml


# --- Tests: apply_translations ---

class TestWordProcessorApplyTranslations:
    """Test WordProcessor.apply_translations()"""

    def test_applies_translations_to_paragraphs(self, processor, sample_docx, tmp_path):
        """Applies translations to paragraphs"""
        output_path = tmp_path / "output.docx"

        # Get block IDs first (only Japanese text is extracted)
        blocks = list(processor.extract_text_blocks(sample_docx))
        assert len(blocks) == 1  # Only Japanese paragraph

        block_0_id = blocks[0].id

        translations = {
            block_0_id: "This is the first paragraph.",
        }

        processor.apply_translations(
            sample_docx, output_path, translations, "jp_to_en"
        )

        # Verify output
        doc = Document(output_path)
        para_texts = [p.text for p in doc.paragraphs if p.text.strip()]

        assert "This is the first paragraph." in para_texts

    def test_applies_translations_to_tables(self, processor, docx_with_table, tmp_path):
        """Applies translations to table cells"""
        output_path = tmp_path / "output.docx"

        translations = {
            "table_0_r0_c0": "Header 1",
            "table_0_r1_c1": "Data",
        }

        processor.apply_translations(
            docx_with_table, output_path, translations, "jp_to_en"
        )

        # Verify output
        doc = Document(output_path)
        table = doc.tables[0]

        assert table.cell(0, 0).text == "Header 1"
        assert table.cell(1, 1).text == "Data"
        # Unchanged
        assert table.cell(0, 1).text == "Header 2"

    def test_preserves_untranslated_content(self, processor, sample_docx, tmp_path):
        """Content not in translations dict is unchanged"""
        output_path = tmp_path / "output.docx"

        blocks = list(processor.extract_text_blocks(sample_docx))
        block_0_id = blocks[0].id

        translations = {
            block_0_id: "Translated first paragraph",
            # Second paragraph not translated
        }

        processor.apply_translations(
            sample_docx, output_path, translations, "jp_to_en"
        )

        doc = Document(output_path)
        para_texts = [p.text for p in doc.paragraphs]

        assert "Translated first paragraph" in para_texts
        assert "This is the second paragraph." in para_texts  # unchanged

    def test_creates_output_file(self, processor, sample_docx, tmp_path):
        """Output file is created"""
        output_path = tmp_path / "output.docx"

        processor.apply_translations(
            sample_docx, output_path, {}, "jp_to_en"
        )

        assert output_path.exists()


# --- Tests: Edge cases ---

class TestWordProcessorEdgeCases:
    """Test edge cases"""

    def test_paragraph_with_multiple_runs(self, processor, tmp_path):
        """Handles paragraphs with multiple runs"""
        file_path = tmp_path / "multi_run.docx"
        doc = Document()

        para = doc.add_paragraph()
        para.add_run("これは")
        para.add_run("複数の")
        para.add_run("ランです")

        doc.save(file_path)

        blocks = list(processor.extract_text_blocks(file_path))
        assert len(blocks) == 1
        assert blocks[0].text == "これは複数のランです"

    def test_nested_tables(self, processor, tmp_path):
        """Handles documents with tables correctly"""
        file_path = tmp_path / "nested.docx"
        doc = Document()

        # Multiple tables
        table1 = doc.add_table(rows=1, cols=1)
        table1.cell(0, 0).text = "テーブル1セル"

        doc.add_paragraph()  # Separator

        table2 = doc.add_table(rows=1, cols=1)
        table2.cell(0, 0).text = "テーブル2セル"

        doc.save(file_path)

        blocks = list(processor.extract_text_blocks(file_path))
        table_blocks = [b for b in blocks if "table" in b.id]

        assert len(table_blocks) == 2
        # Check different table indices
        tables = {b.metadata["table"] for b in table_blocks}
        assert len(tables) == 2

    def test_long_paragraphs(self, processor, tmp_path):
        """Handles long paragraphs"""
        file_path = tmp_path / "long.docx"
        doc = Document()

        long_text = "これは長いテキストです。" * 100
        doc.add_paragraph(long_text)

        doc.save(file_path)

        blocks = list(processor.extract_text_blocks(file_path))
        assert len(blocks) == 1
        assert len(blocks[0].text) > 1000

    def test_unicode_content(self, processor, tmp_path):
        """Handles various Unicode characters"""
        file_path = tmp_path / "unicode.docx"
        doc = Document()

        doc.add_paragraph("日本語テスト")
        doc.add_paragraph("中文测试")  # Chinese uses CJK characters, detected as Japanese
        # Note: Korean and English-only are skipped
        doc.add_paragraph("한국어 테스트")  # Korean (not Japanese)
        doc.add_paragraph("Emoji test 😀🎉")  # English-only

        doc.save(file_path)

        blocks = list(processor.extract_text_blocks(file_path))
        # Only Japanese/CJK containing texts are extracted
        assert len(blocks) == 2


# --- Tests: create_bilingual_document ---

class TestWordProcessorCreateBilingualDocument:
    """Test WordProcessor.create_bilingual_document()"""

    def test_creates_bilingual_document(self, processor, tmp_path):
        """Creates document with original and translated content"""
        # Create original document
        original_path = tmp_path / "original.docx"
        doc_orig = Document()
        doc_orig.add_paragraph("これは日本語の段落です。")
        doc_orig.add_paragraph("これは2番目の段落です。")
        doc_orig.save(original_path)

        # Create translated document
        translated_path = tmp_path / "translated.docx"
        doc_trans = Document()
        doc_trans.add_paragraph("This is a Japanese paragraph.")
        doc_trans.add_paragraph("This is the second paragraph.")
        doc_trans.save(translated_path)

        # Create bilingual document
        output_path = tmp_path / "bilingual.docx"
        result = processor.create_bilingual_document(
            original_path, translated_path, output_path
        )

        # Verify result
        assert result["original_paragraphs"] == 2
        assert result["translated_paragraphs"] == 2
        assert result["total_paragraphs"] == 4

        # Verify output file
        doc_out = Document(output_path)
        para_texts = [p.text for p in doc_out.paragraphs if p.text.strip()]

        # Should contain both original and translated content
        assert "これは日本語の段落です。" in para_texts
        assert "This is a Japanese paragraph." in para_texts

    def test_includes_translation_header(self, processor, tmp_path):
        """Includes 【翻訳】 header before translated content"""
        # Create original
        original_path = tmp_path / "original.docx"
        doc_orig = Document()
        doc_orig.add_paragraph("Original text")
        doc_orig.save(original_path)

        # Create translated
        translated_path = tmp_path / "translated.docx"
        doc_trans = Document()
        doc_trans.add_paragraph("翻訳されたテキスト")
        doc_trans.save(translated_path)

        output_path = tmp_path / "bilingual.docx"
        processor.create_bilingual_document(
            original_path, translated_path, output_path
        )

        doc_out = Document(output_path)
        all_text = "\n".join(p.text for p in doc_out.paragraphs)

        assert "【翻訳】" in all_text

    def test_includes_separator(self, processor, tmp_path):
        """Includes separator line between original and translated"""
        original_path = tmp_path / "original.docx"
        doc_orig = Document()
        doc_orig.add_paragraph("Original")
        doc_orig.save(original_path)

        translated_path = tmp_path / "translated.docx"
        doc_trans = Document()
        doc_trans.add_paragraph("Translated")
        doc_trans.save(translated_path)

        output_path = tmp_path / "bilingual.docx"
        processor.create_bilingual_document(
            original_path, translated_path, output_path
        )

        doc_out = Document(output_path)
        all_text = "\n".join(p.text for p in doc_out.paragraphs)

        # Should contain separator (series of dashes)
        assert "─" in all_text

    def test_handles_empty_documents(self, processor, tmp_path):
        """Handles empty documents gracefully"""
        original_path = tmp_path / "original.docx"
        doc_orig = Document()
        doc_orig.save(original_path)

        translated_path = tmp_path / "translated.docx"
        doc_trans = Document()
        doc_trans.save(translated_path)

        output_path = tmp_path / "bilingual.docx"
        result = processor.create_bilingual_document(
            original_path, translated_path, output_path
        )

        assert result["original_paragraphs"] == 0
        assert result["translated_paragraphs"] == 0
        assert output_path.exists()

    def test_copies_table_content(self, processor, docx_with_table, tmp_path):
        """Copies tables from translated document"""
        # Create translated version with table
        translated_path = tmp_path / "translated.docx"
        doc_trans = Document()
        doc_trans.add_paragraph("Document with table")
        table = doc_trans.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "12345"
        table.cell(1, 1).text = "Data"
        doc_trans.save(translated_path)

        output_path = tmp_path / "bilingual.docx"
        processor.create_bilingual_document(
            docx_with_table, translated_path, output_path
        )

        doc_out = Document(output_path)
        # Should have tables from both documents
        assert len(doc_out.tables) == 2
