# tests/test_base_processor.py
"""
Tests for yakulingo.processors.base - Abstract FileProcessor base class.
Validates interface contract and default implementations.
"""

import pytest
from pathlib import Path
from abc import ABC
from typing import Iterator

from yakulingo.processors.base import FileProcessor
from yakulingo.processors.excel_processor import ExcelProcessor
from yakulingo.processors.word_processor import WordProcessor
from yakulingo.processors.pptx_processor import PptxProcessor
from yakulingo.processors.pdf_processor import PdfProcessor
from yakulingo.models.types import TextBlock, FileInfo, FileType


# =============================================================================
# Tests: Abstract Interface Contract
# =============================================================================

class TestAbstractInterface:
    """Tests for abstract interface contract"""

    def test_file_processor_is_abstract(self):
        """FileProcessor should be an abstract class"""
        assert issubclass(FileProcessor, ABC)

    def test_cannot_instantiate_base_class(self):
        """Cannot instantiate FileProcessor directly"""
        with pytest.raises(TypeError) as exc:
            FileProcessor()

        assert "abstract" in str(exc.value).lower()

    def test_abstract_methods_defined(self):
        """Verify all abstract methods are defined"""
        abstract_methods = {
            'file_type',
            'supported_extensions',
            'get_file_info',
            'extract_text_blocks',
            'apply_translations',
        }

        # Get abstract methods from the class
        processor_abstracts = set(FileProcessor.__abstractmethods__)

        # All expected abstract methods should be present
        for method in abstract_methods:
            assert method in processor_abstracts, f"Missing abstract method: {method}"


# =============================================================================
# Tests: All Subclasses Implement Interface
# =============================================================================

class TestSubclassImplementation:
    """Tests that all subclasses properly implement the interface"""

    @pytest.fixture(params=[
        ExcelProcessor,
        WordProcessor,
        PptxProcessor,
        PdfProcessor,
    ])
    def processor_class(self, request):
        """Parametrized fixture for all processor classes"""
        return request.param

    def test_subclass_can_be_instantiated(self, processor_class):
        """All subclasses can be instantiated"""
        processor = processor_class()
        assert processor is not None

    def test_subclass_has_file_type_property(self, processor_class):
        """All subclasses have file_type property"""
        processor = processor_class()
        assert hasattr(processor, 'file_type')
        assert isinstance(processor.file_type, FileType)

    def test_subclass_has_supported_extensions_property(self, processor_class):
        """All subclasses have supported_extensions property"""
        processor = processor_class()
        assert hasattr(processor, 'supported_extensions')
        extensions = processor.supported_extensions
        assert isinstance(extensions, list)
        assert len(extensions) > 0
        # All extensions should start with .
        for ext in extensions:
            assert ext.startswith('.'), f"Extension should start with '.': {ext}"

    def test_subclass_has_get_file_info_method(self, processor_class):
        """All subclasses have get_file_info method"""
        processor = processor_class()
        assert hasattr(processor, 'get_file_info')
        assert callable(processor.get_file_info)

    def test_subclass_has_extract_text_blocks_method(self, processor_class):
        """All subclasses have extract_text_blocks method"""
        processor = processor_class()
        assert hasattr(processor, 'extract_text_blocks')
        assert callable(processor.extract_text_blocks)

    def test_subclass_has_apply_translations_method(self, processor_class):
        """All subclasses have apply_translations method"""
        processor = processor_class()
        assert hasattr(processor, 'apply_translations')
        assert callable(processor.apply_translations)


# =============================================================================
# Tests: Default should_translate Implementation
# =============================================================================

class TestShouldTranslate:
    """Tests for the default should_translate method"""

    @pytest.fixture
    def processor(self):
        """Get a concrete processor for testing default method"""
        return ExcelProcessor()

    def test_empty_string_not_translated(self, processor):
        """Empty string should not be translated"""
        assert processor.should_translate("") is False

    def test_whitespace_only_not_translated(self, processor):
        """Whitespace-only string should not be translated"""
        assert processor.should_translate("   ") is False
        assert processor.should_translate("\t\n") is False
        assert processor.should_translate("  \n  \t  ") is False

    def test_numbers_only_not_translated(self, processor):
        """Numbers-only string should not be translated"""
        assert processor.should_translate("12345") is False
        assert processor.should_translate("123.45") is False
        assert processor.should_translate("1,234,567") is False
        assert processor.should_translate("-123") is False

    def test_regular_text_translated(self, processor):
        """Regular text should be translated"""
        assert processor.should_translate("Hello") is True
        assert processor.should_translate("こんにちは") is True
        assert processor.should_translate("Hello World") is True

    def test_mixed_text_and_numbers_translated(self, processor):
        """Text with numbers should be translated"""
        assert processor.should_translate("Item 123") is True
        assert processor.should_translate("2024年") is True
        assert processor.should_translate("$100 price") is True

    def test_special_characters_only(self, processor):
        """Special characters handling"""
        # Pure punctuation might be treated differently
        result = processor.should_translate("...")
        # Implementation may vary - just verify it doesn't crash
        assert isinstance(result, bool)


# =============================================================================
# Tests: supports_extension Method
# =============================================================================

class TestSupportsExtension:
    """Tests for the supports_extension method"""

    @pytest.fixture(params=[
        (ExcelProcessor, ['.xlsx', '.xls']),
        (WordProcessor, ['.docx', '.doc']),
        (PptxProcessor, ['.pptx', '.ppt']),
        (PdfProcessor, ['.pdf']),
    ])
    def processor_with_extensions(self, request):
        """Processor with expected extensions"""
        processor_class, expected_extensions = request.param
        return processor_class(), expected_extensions

    def test_supports_own_extensions(self, processor_with_extensions):
        """Processor supports its own extensions"""
        processor, extensions = processor_with_extensions
        for ext in extensions:
            assert processor.supports_extension(ext) is True

    def test_case_insensitive_extension_check(self, processor_with_extensions):
        """Extension check should be case-insensitive"""
        processor, extensions = processor_with_extensions
        for ext in extensions:
            assert processor.supports_extension(ext.upper()) is True
            assert processor.supports_extension(ext.lower()) is True

    def test_does_not_support_other_extensions(self, processor_with_extensions):
        """Processor does not support other extensions"""
        processor, _ = processor_with_extensions
        assert processor.supports_extension('.xyz') is False
        assert processor.supports_extension('.unknown') is False
        assert processor.supports_extension('.txt') is False


# =============================================================================
# Tests: File Type Consistency
# =============================================================================

class TestFileTypeConsistency:
    """Tests for file type consistency across processors"""

    def test_excel_processor_file_type(self):
        """ExcelProcessor has correct file type"""
        processor = ExcelProcessor()
        assert processor.file_type == FileType.EXCEL

    def test_word_processor_file_type(self):
        """WordProcessor has correct file type"""
        processor = WordProcessor()
        assert processor.file_type == FileType.WORD

    def test_pptx_processor_file_type(self):
        """PptxProcessor has correct file type"""
        processor = PptxProcessor()
        assert processor.file_type == FileType.POWERPOINT

    def test_pdf_processor_file_type(self):
        """PdfProcessor has correct file type"""
        processor = PdfProcessor()
        assert processor.file_type == FileType.PDF

    def test_all_file_types_covered(self):
        """All FileType values have a processor"""
        # Collect all processor file types
        processor_types = {
            ExcelProcessor().file_type,
            WordProcessor().file_type,
            PptxProcessor().file_type,
            PdfProcessor().file_type,
        }

        # Check main file types are covered
        main_types = {FileType.EXCEL, FileType.WORD, FileType.POWERPOINT, FileType.PDF}
        assert main_types.issubset(processor_types)


# =============================================================================
# Tests: Extension Mapping
# =============================================================================

class TestExtensionMapping:
    """Tests for extension to processor mapping"""

    def test_xlsx_maps_to_excel(self):
        """xlsx extension maps to ExcelProcessor"""
        processor = ExcelProcessor()
        assert processor.supports_extension('.xlsx')
        assert processor.file_type == FileType.EXCEL

    def test_xls_maps_to_excel(self):
        """xls extension maps to ExcelProcessor"""
        processor = ExcelProcessor()
        assert processor.supports_extension('.xls')

    def test_docx_maps_to_word(self):
        """docx extension maps to WordProcessor"""
        processor = WordProcessor()
        assert processor.supports_extension('.docx')
        assert processor.file_type == FileType.WORD

    def test_doc_maps_to_word(self):
        """doc extension maps to WordProcessor"""
        processor = WordProcessor()
        assert processor.supports_extension('.doc')

    def test_pptx_maps_to_powerpoint(self):
        """pptx extension maps to PptxProcessor"""
        processor = PptxProcessor()
        assert processor.supports_extension('.pptx')
        assert processor.file_type == FileType.POWERPOINT

    def test_ppt_maps_to_powerpoint(self):
        """ppt extension maps to PptxProcessor"""
        processor = PptxProcessor()
        assert processor.supports_extension('.ppt')

    def test_pdf_maps_to_pdf(self):
        """pdf extension maps to PdfProcessor"""
        processor = PdfProcessor()
        assert processor.supports_extension('.pdf')
        assert processor.file_type == FileType.PDF


# =============================================================================
# Tests: Concrete Processor Behavior with Real Files
# =============================================================================

class TestProcessorWithRealFiles:
    """Tests with real file operations"""

    @pytest.fixture
    def sample_excel(self, tmp_path):
        """Create a sample Excel file"""
        import openpyxl
        file_path = tmp_path / "test.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = 'テスト'
        ws['B1'] = 'Test'
        wb.save(file_path)
        return file_path

    @pytest.fixture
    def sample_word(self, tmp_path):
        """Create a sample Word file"""
        from docx import Document
        file_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph('テスト')
        doc.add_paragraph('Test')
        doc.save(file_path)
        return file_path

    def test_excel_get_file_info_returns_file_info(self, sample_excel):
        """ExcelProcessor.get_file_info returns FileInfo"""
        processor = ExcelProcessor()
        info = processor.get_file_info(sample_excel)

        assert isinstance(info, FileInfo)
        assert info.path == sample_excel
        assert info.file_type == FileType.EXCEL
        assert info.size_bytes > 0

    def test_word_get_file_info_returns_file_info(self, sample_word):
        """WordProcessor.get_file_info returns FileInfo"""
        processor = WordProcessor()
        info = processor.get_file_info(sample_word)

        assert isinstance(info, FileInfo)
        assert info.path == sample_word
        assert info.file_type == FileType.WORD
        assert info.size_bytes > 0

    def test_excel_extract_text_blocks_returns_iterator(self, sample_excel):
        """ExcelProcessor.extract_text_blocks returns iterator"""
        processor = ExcelProcessor()
        blocks = processor.extract_text_blocks(sample_excel)

        # Should be iterable
        block_list = list(blocks)
        assert len(block_list) > 0

        # Each block should be TextBlock
        for block in block_list:
            assert isinstance(block, TextBlock)
            assert block.id is not None
            assert block.text is not None

    def test_word_extract_text_blocks_returns_iterator(self, sample_word):
        """WordProcessor.extract_text_blocks returns iterator"""
        processor = WordProcessor()
        blocks = processor.extract_text_blocks(sample_word)

        # Should be iterable
        block_list = list(blocks)
        assert len(block_list) > 0

        # Each block should be TextBlock
        for block in block_list:
            assert isinstance(block, TextBlock)


# =============================================================================
# Tests: Error Handling in Base Methods
# =============================================================================

class TestBaseErrorHandling:
    """Tests for error handling in base methods"""

    def test_should_translate_with_none(self):
        """should_translate handles None gracefully"""
        processor = ExcelProcessor()

        # None should be treated as empty
        # Implementation may raise or return False
        try:
            result = processor.should_translate(None)
            assert result is False
        except (TypeError, AttributeError):
            pass  # Expected if None not handled

    def test_supports_extension_with_none(self):
        """supports_extension handles None gracefully"""
        processor = ExcelProcessor()

        try:
            result = processor.supports_extension(None)
            assert result is False
        except (TypeError, AttributeError):
            pass  # Expected if None not handled

    def test_supports_extension_with_empty_string(self):
        """supports_extension handles empty string"""
        processor = ExcelProcessor()
        result = processor.supports_extension('')
        assert result is False

    def test_supports_extension_without_dot(self):
        """supports_extension handles extension without dot"""
        processor = ExcelProcessor()
        # 'xlsx' without dot should not match '.xlsx'
        result = processor.supports_extension('xlsx')
        # May or may not match depending on implementation
        assert isinstance(result, bool)


# =============================================================================
# Tests: TextBlock Properties
# =============================================================================

class TestTextBlockFromProcessors:
    """Tests for TextBlock objects from processors"""

    @pytest.fixture
    def sample_excel(self, tmp_path):
        """Create a sample Excel file"""
        import openpyxl
        file_path = tmp_path / "test.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws['A1'] = 'First Cell'
        ws['B2'] = 'Second Cell'
        wb.save(file_path)
        return file_path

    def test_text_block_has_unique_ids(self, sample_excel):
        """TextBlocks have unique IDs"""
        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(sample_excel))

        ids = [block.id for block in blocks]
        assert len(ids) == len(set(ids)), "TextBlock IDs should be unique"

    def test_text_block_has_location(self, sample_excel):
        """TextBlocks have location information"""
        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(sample_excel))

        for block in blocks:
            assert block.location is not None
            assert len(block.location) > 0

    def test_text_block_text_not_empty(self, sample_excel):
        """TextBlocks have non-empty text"""
        processor = ExcelProcessor()
        blocks = list(processor.extract_text_blocks(sample_excel))

        for block in blocks:
            assert block.text is not None
            assert len(block.text.strip()) > 0


# =============================================================================
# Tests: Custom Subclass Implementation
# =============================================================================

class TestCustomSubclass:
    """Tests for creating custom subclasses"""

    def test_minimal_valid_subclass(self):
        """Minimal valid subclass can be created"""

        class MinimalProcessor(FileProcessor):
            @property
            def file_type(self) -> FileType:
                return FileType.EXCEL

            @property
            def supported_extensions(self) -> list[str]:
                return ['.test']

            def get_file_info(self, file_path: Path) -> FileInfo:
                return FileInfo(
                    path=file_path,
                    file_type=FileType.EXCEL,
                    size_bytes=0,
                )

            def extract_text_blocks(self, file_path: Path) -> Iterator[TextBlock]:
                return iter([])

            def apply_translations(
                self,
                input_path: Path,
                output_path: Path,
                translations: dict[str, str],
                direction: str = "jp_to_en",
            ) -> None:
                pass

        processor = MinimalProcessor()
        assert processor.file_type == FileType.EXCEL
        assert processor.supported_extensions == ['.test']

    def test_incomplete_subclass_fails(self):
        """Incomplete subclass cannot be instantiated"""

        class IncompleteProcessor(FileProcessor):
            @property
            def file_type(self) -> FileType:
                return FileType.EXCEL
            # Missing other abstract methods

        with pytest.raises(TypeError):
            IncompleteProcessor()
